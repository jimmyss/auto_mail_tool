#!/usr/bin/python
# -*- coding: UTF-8 -*-
import os
import time
import tkinter as tk
from re import match
from tkinter import filedialog
import ttkbootstrap as ttk
import win32com.client as win32
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import load_workbook
from ttkbootstrap.constants import *

excel_path=""
word_path=""

class outlook():
    def init(self):
        pass

    def openoutlook(self):
        pass

    def sendmail(self, receivers, cc_receivers, title, body, attach_path=None):
        '''
        发送邮件
        ：param receivers：收件人
        ：param title：主题
        ：param body：邮件内容
        ：param attach_path：附件地址
        ；retuxn：发送
        '''
        outlook = win32.Dispatch('Outlook.Application')
        mail = outlook.CreateItem(0)

        if isinstance(receivers, list):
            if len(receivers) > 1:
                mail.To = ';'.join(receivers)
            else:
                mail.To = receivers[0]
        else:
            mail.To = receivers
        if isinstance(cc_receivers, list):#多个人
            if len(cc_receivers) > 1:
                mail.Cc = ';'.join(cc_receivers)
            else:
                mail.Cc = cc_receivers[0]
        else:#单个人
            mail.Cc = cc_receivers
        mail.Subject = title
        mail.Body = body
        if attach_path is None:
            pass
        else:
            mail.Attachments.Add(attach_path)
        mail.Send()

    def draftmail(self, receivers, cc_receivers, title, body, attach_path=None):
        '''
        保存一份草稿
        :param receivera:收件人
        :param title:主题
        :param body:邮件内容
        :param attach_path:附件地址
        :return:保存
        '''
        outlook = win32.Dispatch('Outlook.Application')
        mail = outlook.CreateItem(0)

        # 如果多个人,这里做个判断
        if isinstance(receivers, list):#多个人
            if len(receivers) > 1:
                mail.To = ';'.join(receivers)
            else:
                mail.To = receivers[0]
        else:#单个人
            mail.To = receivers
        #同样对抄送人也做处理
        if isinstance(cc_receivers, list):#多个人
            if len(cc_receivers) > 1:
                mail.Cc = ';'.join(cc_receivers)
            else:
                mail.Cc = cc_receivers[0]
        else:#单个人
            mail.Cc = cc_receivers

        mail.Subject = title
        mail.Body = body
        if attach_path is not None:
            mail.Attachments.Add(attach_path)
        mail.Save()


def set_text(text):
    message_box.config(state=NORMAL)
    message_box.delete("1.0","end")
    message_box.insert(END,'\n'+str(text))
    message_box.config(state=DISABLED)

def replace(obj):
    if obj is None:
        obj = ''
        return obj

def send_mail(mail_subject, mail_content, xlsx_file_path, docx_file_path, type):
    path=xlsx_file_path.rsplit("/", 1)[0]
    try:
        wb = load_workbook(xlsx_file_path, data_only=True)  # 用openpyxl加载excel表格
        ws = wb.active  # 读取Sheet1内容
        mail_sendto_list = []  # 保存收件人列表
        mail_copy_list = []  # 保存抄送人列表
        col_names = []
        tuples = []
        row_max = ws.max_row  # 保存最大行数
        col_max = ws.max_column  # 保存最大列数
        col_copy = []  # 保存抄送人列号
        start_row = 0
        mail_col_position = 0
        end_flag = False
        for i in range(row_max):
            for j in range(col_max):
                if ws.cell(row=i + 1, column=j + 1).value:
                    # 识别关键字"邮箱"，确定读取开始行
                    if match(r"邮箱", str(ws.cell(row=i + 1, column=j + 1).value)):
                        start_row = i + 1
                        mail_col_position = j + 1
                        # 识别关键字“抄送人”，确定有几个抄送人
                        for k in range(j + 1, col_max):
                            if match(r"邮箱", ws.cell(row=i + 1, column=k + 1).value):  # 如果有邮箱的字眼，就默认为抄送人
                                col_copy.append(k + 1)
                        end_flag = True
                        break
            if end_flag:
                break
        for i in range(start_row, row_max + 1):
            if i == start_row:
                # 确定要读取的列名
                for j in range(col_max):
                    col_names.append((ws.cell(row=i, column=j + 1).value).replace("\n", ""))  # 去除列里面的换行符
            else:
                # 读取每一行
                if ws.cell(row=i, column=mail_col_position).value:
                    tuples.append(ws[i])  # 保存每行元组
                    # 填写收件人
                    mail_sendto_list.append(ws.cell(row=i, column=mail_col_position).value)
                    # 填写抄送人
                    temp_copy_mail = []
                    for j in range(len(col_copy)):
                        # 若有抄送人，将所有抄送人保存到一个list中并保存到mail_copy_list中
                        if ws.cell(row=i, column=col_copy[j]).value:
                            temp_copy_mail.append(ws.cell(row=i, column=col_copy[j]).value)
                    mail_copy_list.append(temp_copy_mail)
    except Exception as n:
        return "无法打开或读取excel，请检查excel文件名和setting中待发送excel文件全名是否一致，请把excel放在程序文件夹下再运行"


    # -----将excel提取到的数据填入表格------
    try:
        for context in tuples:
            try:
                # 打开待操作的docx文件，r表示“防止\转义”
                document = Document(docx_file_path)
            except Exception as e:
                return "不能打开doc格式，请手动转换模板为docx再执行程序"
            # 提取文件中的表格
            tables = document.tables

            # 循环填写表格
            for table in tables:
                for i in range(0, len(table.rows)):
                    for j in range(0, len(table.columns)):
                        table_col_name = table.cell(i, j).text.replace("\n", "")
                        if table_col_name in col_names:  # 若检测到对应列
                            try:
                                if context[col_names.index(table_col_name)].value != None:
                                    # 若非空，将该列下面填入值
                                    table.cell(i + 1, j).text = str(context[col_names.index(table_col_name)].value)
                                    table.cell(i + 1, j).paragraphs[
                                        0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 居中显示
                                else:
                                    # 若为空值，填入空白
                                    table.cell(i + 1, j).text = ""
                            except Exception as e:
                                return "word表格格式错误，应该每张表仅有两行"

            # -------替换word文档中的姓名/工号姓名-------
            if "姓名" in col_names:
                name = context[col_names.index("姓名")].value
            if "工号" in col_names:
                number = context[col_names.index("工号")].value

            start_flag = False  # 识别是否开始输入的
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if "XXX" in run.text:
                        run.text = run.text.replace('XXX', name)
                    if "YYY" in run.text:
                        run.text = run.text.replace('YYY', number)

            # -------保存至新文件，文件名自拟-----
            docx_file_name=docx_file_path.split('/')[len(docx_file_path.split('/'))-1]
            document.save(r'./' + name + "-" + docx_file_name)
    except Exception as n:
        return "生成word失败，请检查word文件名与setting中word模板文件全名是否一致，请把word放在程序文件夹下再运行"

    time.sleep(3)
    #按照用户选择是保存到草稿箱还是保存到发件箱来处理邮件
    otlk = outlook()
    count=0 # 对收件人计数
    for context in tuples: #将每个文件保存到草稿箱
        # 设置收件人邮箱和抄送人邮箱，可以为多个收件人
        mail_receiver = []#收件人
        mail_receiver.append(mail_sendto_list[count])
        mail_cc_receiver = []#抄送人
        for copyer in mail_copy_list[count]:
            mail_cc_receiver.append(copyer)
        if "姓名" in col_names:#文件收件人姓名
            name = context[col_names.index("姓名")].value
        #attachment=r'D:\University\notes\北京交通大学软件学院大二杜世茂.pdf'
        attachment= path +"/" + name + "-" + docx_file_name
        if type == 1:  # 保存到草稿箱
            otlk.draftmail(title=mail_subject, body=mail_content, receivers=mail_receiver, cc_receivers=mail_cc_receiver, attach_path=attachment)
        elif type==0: #直接发送
            otlk.sendmail(title=mail_subject, body=mail_content, receivers=mail_receiver, cc_receivers=mail_cc_receiver, attach_path=attachment)
        time.sleep(3)
    if type==1:
        return "所有邮件已经保存到草稿箱，请前往Outlook查看。"
    elif type==0:
        return "所有邮件已经自动发送，若Outlook处于脱机状态，请前往发件箱查看。"

#按钮事件
def press_send(type):
    global excel_path
    global word_path
    set_text("")
    #获取输入框信息
    title_=title.get().split('\n')[0]
    text_=text.get().split('\n')[0]
    excel_=excel_path
    word_=word_path
    #检测输入框是否都有信息
    if title_ and text_ and excel_ and word_:
        #若都有信息，发送邮件
        result=send_mail(title_, text_, excel_,word_, type=type)
        set_text(result)
    else:
        message_box.config(state=NORMAL)
        message_box.insert(END,'请输入所有信息再点击发送')
        message_box.config(state=DISABLED)

def open_file(num):
    global excel_path
    global word_path
    file_path = filedialog.askopenfilename(title=u'选择文件', initialdir=(os.path.expanduser('H:/')))
    if file_path is not None:
        if num==0:#打开setting.txt文件
            with open(file=file_path, mode='r+', encoding='utf-8') as file:
                title.delete(0, END)
                text.delete(0, END)
                excel.delete(0, END)
                word.delete(0, END)
                t = file.readline().split(':')[1]
                title.insert(0, t)
                t = file.readline().split(':')[1]
                text.insert(0, t)
                t = file.readline().split(':',1)[1].replace("\n","")
                excel_path=t
                excel.insert(0, t.split("\\")[len(t.split("\\"))-1])
                t = file.readline().split(':',1)[1].replace("\n","")
                word_path=t
                word.insert(0, t.split("\\")[len(t.split("\\"))-1])
        elif num==1:
            excel.delete(0,END)
            excel_path=file_path
            excel.insert(0,file_path.split('/')[len(file_path.split('/'))-1])
        elif num==2:
            word.delete(0,END)
            word_path=file_path
            word.insert(0,file_path.split('/')[len(file_path.split('/'))-1])


# 调用Tk()创建主窗口
root_window =tk.Tk()
# 给主窗口起一个名字，也就是窗口的名字
root_window.title('自动办公工具')
# 设置窗口大小:宽x高,注,此处不能为 "*",必须使用 "x"
root_window.geometry('550x550')

inputFrame=ttk.Frame(root_window,padding=(10,5,10,0))


title_l=ttk.Label(inputFrame,text="邮件主题:")
title_l.grid(row=3,sticky=E,pady=10)
text_l=ttk.Label(inputFrame,text="邮件正文内容:")
text_l.grid(row=4,sticky=E,pady=10)
excel_l=ttk.Label(inputFrame,text="待发送excel文件全名:")
excel_l.grid(row=5,sticky=E,pady=10)
word_l=ttk.Label(inputFrame,text="word模板文件全名:")
word_l.grid(row=6,sticky=E,pady=10)
result_l=ttk.Label(inputFrame, text="结果显示:")
result_l.grid(row=8,sticky=NE,pady=10)

#输入控件
title=ttk.Entry(inputFrame, width=35)
title.grid(row=3,column=1)
text=ttk.Entry(inputFrame, width=35)
text.grid(row=4,column=1)
excel=ttk.Entry(inputFrame, width=35)
excel.grid(row=5,column=1)
word=ttk.Entry(inputFrame, width=35)
word.grid(row=6,column=1)

#按钮控件
send_button=ttk.Button(inputFrame, text="直接发送", command=lambda: press_send(0))
send_button.grid(row=7, column=1,pady=10)
open_file_button=ttk.Button(inputFrame, text="信息自动填入", command=lambda: open_file(0))
open_file_button.grid(row=7, column=0)
select_excel=ttk.Button(inputFrame, text="选择excel", command=lambda: open_file(1))
select_excel.grid(row=5, column=2, sticky=W,padx=10)
select_word=ttk.Button(inputFrame, text="选择word", command=lambda: open_file(2))
select_word.grid(row=6, column=2, sticky=W,padx=10)
savedraft_button=ttk.Button(inputFrame, text="保存到草稿箱", command=lambda: press_send(1))
savedraft_button.grid(row=7, column=2, pady=10)

#消息框控件
message_box=ttk.Text(inputFrame, height=10, width=40)
message_box.config(state=DISABLED)
message_box.grid(row=8,column=1,columnspan=8, sticky=W)

inputFrame.grid(row = 0, column = 0, columnspan=3)

#开启主循环，让窗口处于显示状态
root_window.mainloop()